import sys
import json
import os
import subprocess
import math
import random

from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
                             QLabel, QLineEdit, QPushButton, QTextEdit, QComboBox, 
                             QFileDialog, QMessageBox, QFrame, QTabWidget, QSplitter, 
                             QListWidget, QListWidgetItem, QShortcut, QSlider, QStyleFactory)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QTimer, QPointF
from PyQt5.QtGui import QColor, QFont, QTextCursor, QKeySequence, QTextCharFormat, QPainter, QPen, QBrush

# --- 尝试导入外部依赖 ---
try:
    from ai_visual_window import AIVisualWindow
    HAS_VISUAL_WINDOW = True
except ImportError:
    HAS_VISUAL_WINDOW = False
    class AIVisualWindow(QWidget):
        def add_stream_char(self, char): pass

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

CONFIG_FILE = "gui_configs.json"

# === 全局统一样式表 (基础) ===
# 注意：为了支持透明度调节，部分背景色将在代码中动态生成
GLOBAL_STYLESHEET = """
QMainWindow {
    background-color: #0d1117;
    color: #c9d1d9;
}
QWidget {
    color: #c9d1d9;
}
QTabWidget::pane {
    border: 1px solid #30363d;
    background: transparent; /* 让Tab内容决定背景 */
}
QTabBar::tab {
    background: #161b22;
    color: #8b949e;
    padding: 8px 20px;
    border: 1px solid #30363d;
    border-bottom: none;
    border-top-left-radius: 4px;
    border-top-right-radius: 4px;
    margin-right: 2px;
}
QTabBar::tab:selected {
    background: #0d1117;
    color: #58a6ff;
    border-bottom: 2px solid #58a6ff;
}
QTabBar::tab:hover {
    background: #21262d;
}
QLabel {
    color: #58a6ff; 
    font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;
    font-weight: bold;
}
QLineEdit {
    background-color: #161b22;
    border: 1px solid #30363d;
    border-radius: 4px;
    color: #c9d1d9;
    padding: 6px;
    font-family: 'Consolas', 'Microsoft YaHei';
}
QLineEdit:focus {
    border: 1px solid #58a6ff;
    background-color: #0d1117;
}
QPushButton {
    background-color: #238636;
    color: white;
    border: none;
    padding: 6px 12px;
    border-radius: 6px;
    font-weight: bold;
    font-size: 13px;
}
QPushButton:hover {
    background-color: #2ea043;
}
QPushButton:pressed {
    background-color: #1a6329;
}
QPushButton#VisualBtn {
    background-color: #1f6feb;
    border: 1px solid #1f6feb;
}
QPushButton#VisualBtn:hover {
    background-color: #388bfd;
}
QTextEdit {
    background-color: #0d1117;
    border: 1px solid #30363d;
    color: #c9d1d9; 
    font-family: 'Consolas', 'Microsoft YaHei', monospace;
    font-size: 13px;
    line-height: 1.5;
}
QComboBox {
    background-color: #161b22;
    color: #c9d1d9;
    border: 1px solid #30363d;
    padding: 5px;
    border-radius: 4px;
}
QComboBox::drop-down {
    border: none;
    background: transparent;
}
QComboBox QAbstractItemView {
    background-color: #161b22;
    color: #c9d1d9;
    selection-background-color: #1f6feb;
    border: 1px solid #30363d;
}
QListWidget { 
    background-color: #0d1117; 
    border: 1px solid #30363d; 
    border-radius: 6px;
    color: #c9d1d9; 
    font-size: 14px; 
    padding: 5px;
}
QListWidget::item { padding: 5px; }
QListWidget::item:selected { background-color: #1f6feb; border-radius: 4px; color: white; }
QSplitter::handle { background-color: #30363d; }
QSlider::groove:horizontal {
    border: 1px solid #30363d;
    height: 4px;
    background: #161b22;
    margin: 2px 0;
    border-radius: 2px;
}
QSlider::handle:horizontal {
    background: #58a6ff;
    border: 1px solid #58a6ff;
    width: 14px;
    height: 14px;
    margin: -5px 0;
    border-radius: 7px;
}
"""

# =================================================================================
# 模块 0: 粒子向量视觉特效 (兼容 Win7/Py3.8)
# =================================================================================

class NeuralParticleOverlay(QWidget):
    """
    全屏透明覆盖层，绘制粒子和连接线，模拟神经网络或星空效果。
    """
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAttribute(Qt.WA_TransparentForMouseEvents) # 鼠标穿透
        self.setAttribute(Qt.WA_NoSystemBackground)        # 无背景
        
        self.particles = []
        self.num_particles = 60  # 稍微增加粒子数量
        self.connect_distance = 140 # 连线距离阈值
        self.timer = QTimer(self)
        self.timer.timeout.connect(self.update_particles)
        self.timer.start(40) # 25 FPS
        
        self.initialized = False

    def init_particles(self):
        self.particles = []
        w = self.width()
        h = self.height()
        for _ in range(self.num_particles):
            self.particles.append({
                'x': random.uniform(0, w),
                'y': random.uniform(0, h),
                'vx': random.uniform(-0.6, 0.6), 
                'vy': random.uniform(-0.6, 0.6), 
                'size': random.uniform(2, 4.5)
            })
        self.initialized = True

    def resizeEvent(self, event):
        if not self.initialized:
            self.init_particles()
        super().resizeEvent(event)

    def update_particles(self):
        w = self.width()
        h = self.height()
        
        for p in self.particles:
            p['x'] += p['vx']
            p['y'] += p['vy']
            
            # 边界反弹
            if p['x'] < 0 or p['x'] > w: p['vx'] *= -1
            if p['y'] < 0 or p['y'] > h: p['vy'] *= -1
            
        self.update() 

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)
        
        # 颜色：青色系，低透明度
        particle_color = QColor(88, 166, 255, 60) 
        line_color = QColor(88, 166, 255, 30)
        
        brush = QBrush(particle_color)
        pen_particle = QPen(Qt.NoPen)
        pen_line = QPen(line_color)
        pen_line.setWidth(1)
        
        points = []
        for p in self.particles:
            pt = QPointF(p['x'], p['y'])
            points.append(pt)
            
            painter.setBrush(brush)
            painter.setPen(pen_particle)
            painter.drawEllipse(pt, p['size'], p['size'])
            
        # 绘制连线
        painter.setPen(pen_line)
        for i in range(len(points)):
            for j in range(i + 1, len(points)):
                # 简单距离计算
                dx = points[i].x() - points[j].x()
                dy = points[i].y() - points[j].y()
                dist_sq = dx*dx + dy*dy
                
                if dist_sq < self.connect_distance * self.connect_distance:
                    painter.drawLine(points[i], points[j])


# =================================================================================
# 模块 1: 索引构建 Tab
# =================================================================================

class WorkerThread(QThread):
    log_signal = pyqtSignal(str)      
    stream_signal = pyqtSignal(str)   

    def __init__(self, command):
        super().__init__()
        self.command = command
        self.line_buffer = ""

    def run(self):
        process = subprocess.Popen(
            self.command,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT, 
            shell=True,
            text=True,
            encoding='utf-8',
            errors='replace',
            bufsize=0 
        )

        while True:
            char = process.stdout.read(1)
            if not char and process.poll() is not None:
                break
            if char:
                self.process_char(char)
        
        self.flush_buffer()
        process.wait()

    def flush_buffer(self):
        if self.line_buffer:
            line = self.line_buffer.strip()
            if line:
                self.emit_log_line(line)
            self.line_buffer = ""

    def process_char(self, char):
        self.line_buffer += char
        if char == "\n":
            line = self.line_buffer.strip()
            if line: 
                if line.startswith("DEBUG_AI_CHAR:"):
                    try:
                        content = line.split("DEBUG_AI_CHAR:", 1)[1]
                        self.stream_signal.emit(content)
                    except: pass
                else:
                    self.emit_log_line(line)
            self.line_buffer = ""

    def emit_log_line(self, line):
        if "[SUCCESS]" in line:
            formatted_line = f"<span style='color:#00FF00; font-weight:bold; font-size:13px;'>{line}</span>"
        elif "[ERROR]" in line or "Exception" in line:
            formatted_line = f"<span style='color:#FF3333; font-weight:bold;'>{line}</span>"
        elif "[INFO]" in line:
            formatted_line = f"<span style='color:#33CCFF;'>{line}</span>"
        elif "[Warning]" in line:
            formatted_line = f"<span style='color:#FFFF00;'>{line}</span>"
        else:
            formatted_line = line
        self.log_signal.emit(formatted_line)

class IndexerTab(QWidget): 
    def __init__(self):
        super().__init__()
        self.setObjectName("IndexerTab")
        if HAS_VISUAL_WINDOW:
            self.visual_window = AIVisualWindow()
        else:
            self.visual_window = QWidget()
        
        # 定义可用模型列表
        self.available_models = [
            "DeepSeek-V3", 
            "DeepSeek-R1", 
            "qwq-32b", 
            "Qwen2.5-32B", 
            "qwen2.5-vl-72b", 
            "xinghuo-x1", 
            "xinghuo-x1-think"
        ]
        
        self.configs = self.load_configs()
        self.init_ui()
        
        # 初始化默认透明度
        self.update_glass_effect(95) # 默认基本不透明

    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 10)

        # Header
        header_layout = QHBoxLayout()
        title_label = QLabel("INDEXER & PROCESSOR")
        title_label.setStyleSheet("font-size: 20px; color: #00ffcc; letter-spacing: 2px;")
        header_layout.addWidget(title_label)
        header_layout.addStretch()
        layout.addLayout(header_layout)

        # Config Frame
        self.cfg_frame = QFrame()
        self.cfg_frame.setObjectName("ConfigFrame")
        cfg_layout = QHBoxLayout(self.cfg_frame)
        
        self.cb_configs = QComboBox()
        self.cb_configs.addItems(self.configs.keys())
        self.cb_configs.currentTextChanged.connect(self.load_selected_config)
        
        btn_save = QPushButton("💾 SAVE CONFIG")
        btn_save.clicked.connect(self.save_config)
        btn_save.setStyleSheet("background-color: #21262d; border: 1px solid #30363d;")

        cfg_layout.addWidget(QLabel("CONFIGURATION:"))
        cfg_layout.addWidget(self.cb_configs, 1)
        cfg_layout.addWidget(btn_save)
        layout.addWidget(self.cfg_frame)

        # Input Area
        input_layout = QVBoxLayout()
        
        # File Select
        file_layout = QHBoxLayout()
        self.edit_pdf = QLineEdit()
        self.edit_pdf.setPlaceholderText("Select PDF document path...")
        btn_file = QPushButton("📂 BROWSE")
        btn_file.clicked.connect(self.get_file)
        file_layout.addWidget(QLabel("DOCUMENT:"))
        file_layout.addWidget(self.edit_pdf, 1)
        file_layout.addWidget(btn_file)
        input_layout.addLayout(file_layout)
        
        # Model Select (Combobox)
        model_layout = QHBoxLayout()
        self.combo_model = QComboBox()
        self.combo_model.addItems(self.available_models)
        self.combo_model.setCurrentText("DeepSeek-V3") # 默认值
        
        model_layout.addWidget(QLabel("AI MODEL:"))
        model_layout.addWidget(self.combo_model, 1)
        input_layout.addLayout(model_layout)
        
        layout.addLayout(input_layout)

        # Buttons
        btn_layout = QHBoxLayout()
        self.btn_run = QPushButton("🚀 INITIALIZE INDEXING")
        self.btn_run.setFixedHeight(45)
        self.btn_run.clicked.connect(self.start_task)
        
        self.btn_visual = QPushButton("👁️ VISUALIZER: OFF")
        self.btn_visual.setObjectName("VisualBtn")
        self.btn_visual.setCheckable(True)
        self.btn_visual.setFixedHeight(45)
        self.btn_visual.clicked.connect(self.toggle_visual_window)
        if not HAS_VISUAL_WINDOW:
             self.btn_visual.setEnabled(False)
             self.btn_visual.setText("👁️ VISUALIZER (Missing)")

        btn_layout.addWidget(self.btn_run, 2)
        btn_layout.addWidget(self.btn_visual, 1)
        layout.addLayout(btn_layout)

        # Console
        layout.addWidget(QLabel("SYSTEM LOGS:"))
        self.txt_console = QTextEdit()
        self.txt_console.setReadOnly(True)
        # 初始样式
        self.txt_console.setStyleSheet("color: #00ff99; font-family: 'Consolas', monospace; font-size: 12px;")
        layout.addWidget(self.txt_console)
        
        # --- Bottom Control Bar (Transparency) ---
        bottom_bar = QWidget()
        bottom_bar.setFixedHeight(30)
        bot_layout = QHBoxLayout(bottom_bar)
        bot_layout.setContentsMargins(0, 0, 0, 0)
        
        lbl_trans = QLabel("✨ UI透明度 (Glass Effect):")
        lbl_trans.setStyleSheet("font-weight: normal; font-size: 12px; color: #8b949e;")
        
        self.slider_alpha = QSlider(Qt.Horizontal)
        self.slider_alpha.setRange(20, 100) # 20% to 100% opacity
        self.slider_alpha.setValue(95)
        self.slider_alpha.setFixedWidth(200)
        self.slider_alpha.valueChanged.connect(self.update_glass_effect)
        
        bot_layout.addStretch()
        bot_layout.addWidget(lbl_trans)
        bot_layout.addWidget(self.slider_alpha)
        
        layout.addWidget(bottom_bar)

    def update_glass_effect(self, value):
        """
        更新UI组件的背景透明度，以突出背后的粒子效果
        value: 0-100 (opacity)
        """
        alpha_hex = f"{int(value * 2.55):02x}" # Convert 0-100 to 00-ff
        
        # 1. Config Frame Background
        self.cfg_frame.setStyleSheet(f"""
            QFrame#ConfigFrame {{
                background-color: #161b22{alpha_hex}; 
                border-radius: 8px; 
                border: 1px solid #30363d;
                padding: 10px;
            }}
        """)
        
        # 2. Text Console Background
        self.txt_console.setStyleSheet(f"""
            QTextEdit {{
                background-color: #0d1117{alpha_hex};
                border: 1px solid #30363d;
                color: #00ff99; 
                font-family: 'Consolas', monospace; 
                font-size: 12px;
            }}
        """)
        
        # 3. Input Fields & Combos
        common_style = f"background-color: #161b22{alpha_hex}; border: 1px solid #30363d; color: #c9d1d9;"
        self.edit_pdf.setStyleSheet(common_style)
        self.combo_model.setStyleSheet(f"""
            QComboBox {{ {common_style} padding: 5px; border-radius: 4px; }}
            QComboBox::drop-down {{ border: none; }}
            QComboBox QAbstractItemView {{ background-color: #161b22; color: #c9d1d9; }} 
        """)
        # 注意: QAbstractItemView 下拉列表通常保持不透明，否则很难看清文字

    def toggle_visual_window(self):
        if not HAS_VISUAL_WINDOW: return
        if self.btn_visual.isChecked():
            self.visual_window.show()
            self.btn_visual.setText("👁️ VISUALIZER: ON")
            window = self.window()
            if window:
                geo = window.geometry()
                self.visual_window.move(geo.x() + geo.width() + 10, geo.y())
        else:
            self.visual_window.hide()
            self.btn_visual.setText("👁️ VISUALIZER: OFF")

    def load_configs(self):
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, 'r', encoding='utf-8') as f: return json.load(f)
            except: pass
        return {"Default": {"pdf": "", "model": "DeepSeek-V3", "pages": "3"}}

    def save_config(self):
        name = self.cb_configs.currentText() or "NewConfig"
        self.configs[name] = {"pdf": self.edit_pdf.text(), "model": self.combo_model.currentText(), "pages": "3"}
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f: json.dump(self.configs, f)
        QMessageBox.information(self, "System", "Configuration Saved Successfully.")

    def load_selected_config(self, name):
        if name in self.configs:
            c = self.configs[name]
            self.edit_pdf.setText(c.get('pdf',''))
            model = c.get('model', 'DeepSeek-V3')
            # 如果配置中的模型不在列表中，添加到列表
            if self.combo_model.findText(model) == -1:
                self.combo_model.addItem(model)
            self.combo_model.setCurrentText(model)

    def get_file(self):
        f, _ = QFileDialog.getOpenFileName(self, "Select PDF", "", "*.pdf")
        if f: self.edit_pdf.setText(f)

    def append_log(self, text):
        self.txt_console.append(text)
        cursor = self.txt_console.textCursor()
        cursor.movePosition(QTextCursor.End)
        self.txt_console.setTextCursor(cursor)

    def start_task(self):
        pdf_path = self.edit_pdf.text()
        if not pdf_path:
            QMessageBox.warning(self, "Error", "Please select a PDF file first.")
            return

        py_exe = sys.executable
        # 使用 combox 的文本
        model_name = self.combo_model.currentText()
        cmd = f'"{py_exe}" -u run_pageindex.py --pdf_path "{pdf_path}" --model "{model_name}" --toc-check-pages 3'
        
        self.txt_console.clear()
        self.txt_console.append(f"<span style='color:#FFFF00'>[SYSTEM] Initializing subprocess with model: {model_name}...</span>")
        
        self.worker = WorkerThread(cmd)
        self.worker.log_signal.connect(self.append_log)
        
        if HAS_VISUAL_WINDOW:
            self.worker.stream_signal.connect(self.visual_window.add_stream_char)
            if not self.btn_visual.isChecked():
                self.btn_visual.click()
            
        self.worker.start()


# =================================================================================
# 模块 2: 知识召回 Tab
# =================================================================================

class RecallTab(QWidget): 
    def __init__(self):
        super().__init__()
        self.data = None
        self.all_nodes = [] 
        self.last_loaded_path = None
        
        self.init_ui()
        self.setup_shortcuts()
        
        # 初始化默认值
        self.current_alpha = 255
        self.current_font_size = 30
        self.change_font_size(self.slider_font.value())
        self.update_transparency(95)

    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(10, 10, 10, 0) # Bottom 0 to fit bar closely
        layout.setSpacing(5)

        # --- Top Bar ---
        top_bar = QHBoxLayout()
        
        self.btn_load = QPushButton("📂 加载索引")
        self.btn_load.clicked.connect(self.open_file_dialog)
        
        self.btn_refresh = QPushButton("🔄 刷新")
        self.btn_refresh.setToolTip("重新加载当前文件并显示全部内容")
        self.btn_refresh.clicked.connect(self.refresh_data)
        
        self.edit_search = QLineEdit()
        self.edit_search.setPlaceholderText("🔍 输入关键词进行全局内容召回...")
        self.edit_search.returnPressed.connect(self.search_content)
        
        self.btn_search = QPushButton("执行召回")
        self.btn_search.clicked.connect(self.search_content)

        self.combo_export = QComboBox()
        self.combo_export.addItems(["DOCX (Word)", "TXT (文本)", "CSV (表格)", "XLSX (Excel)"])
        self.combo_export.setFixedWidth(120)
        
        self.btn_export = QPushButton("💾 导出全文")
        self.btn_export.clicked.connect(self.export_data)
        
        top_bar.addWidget(self.btn_load)
        top_bar.addWidget(self.btn_refresh)
        top_bar.addWidget(self.edit_search, 4)
        top_bar.addWidget(self.btn_search)
        top_bar.addSpacing(20)
        top_bar.addWidget(QLabel("格式:"))
        top_bar.addWidget(self.combo_export)
        top_bar.addWidget(self.btn_export)
        
        layout.addLayout(top_bar)

        # --- Splitter ---
        splitter = QSplitter(Qt.Horizontal)
        
        # Left: Results
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.addWidget(QLabel("召回结果列表:"))
        self.list_results = QListWidget()
        self.list_results.itemClicked.connect(self.display_node_detail)
        left_layout.addWidget(self.list_results)
        
        # Right: Details
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.setContentsMargins(0, 0, 0, 0)
        right_layout.setSpacing(5)
        
        right_layout.addWidget(QLabel("详情预览:"))
        self.txt_header = QTextEdit()
        self.txt_header.setReadOnly(True)
        self.txt_header.setMaximumHeight(150)
        self.txt_header.setStyleSheet("border: none; background-color: #0d1117;") 
        right_layout.addWidget(self.txt_header)

        # Detail Search
        search_bar_layout = QHBoxLayout()
        search_label = QLabel("🔎 正文查找:")
        search_label.setStyleSheet("color: #8b949e; font-size: 12px;")
        
        self.edit_inner_search = QLineEdit()
        self.edit_inner_search.setPlaceholderText("在此处输入文本，按回车高亮显示 (Ctrl+F)")
        self.edit_inner_search.textChanged.connect(self.highlight_text_in_detail)
        self.edit_inner_search.setStyleSheet("""
            background-color: #21262d; border: 1px solid #30363d; 
            color: #ffd700; font-weight: bold;
        """)
        
        search_bar_layout.addWidget(search_label)
        search_bar_layout.addWidget(self.edit_inner_search)
        right_layout.addLayout(search_bar_layout)

        self.txt_content = QTextEdit()
        self.txt_content.setReadOnly(True)
        right_layout.addWidget(self.txt_content)
        
        splitter.addWidget(left_widget)
        splitter.addWidget(right_widget)
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 3)
        
        layout.addWidget(splitter, 1)

        # --- Bottom Control Bar (Font & Transparency) ---
        bottom_bar = QWidget()
        bottom_bar.setFixedHeight(36)
        
        font_layout = QHBoxLayout(bottom_bar)
        font_layout.setContentsMargins(10, 0, 10, 0) 
        font_layout.setSpacing(15)
        
        # Font Slider
        lbl_font_icon = QLabel("🔠 字号调节:")
        lbl_font_icon.setStyleSheet("color: #c9d1d9; font-weight: normal; font-size: 12px;")
        
        self.slider_font = QSlider(Qt.Horizontal)
        self.slider_font.setRange(12, 40)
        self.slider_font.setValue(30)
        self.slider_font.setFixedWidth(150)
        self.slider_font.valueChanged.connect(self.change_font_size)
        
        self.lbl_font_val = QLabel("30px")
        self.lbl_font_val.setStyleSheet("color: #58a6ff; font-weight: bold; min-width: 40px; font-size: 12px;")

        # Transparency Slider
        lbl_trans_icon = QLabel("✨ UI透明度:")
        lbl_trans_icon.setStyleSheet("color: #c9d1d9; font-weight: normal; font-size: 12px;")

        self.slider_trans = QSlider(Qt.Horizontal)
        self.slider_trans.setRange(20, 100)
        self.slider_trans.setValue(95)
        self.slider_trans.setFixedWidth(150)
        self.slider_trans.valueChanged.connect(self.update_transparency)

        font_layout.addStretch()
        font_layout.addWidget(lbl_trans_icon)
        font_layout.addWidget(self.slider_trans)
        font_layout.addWidget(self.lbl_font_val) # Spacer
        font_layout.addWidget(lbl_font_icon)
        font_layout.addWidget(self.slider_font)
        font_layout.addWidget(self.lbl_font_val)
        
        layout.addWidget(bottom_bar, 0)

    def update_transparency(self, val):
        self.current_alpha = int(val * 2.55)
        # 刷新当前样式，应用新的透明度
        self.change_font_size(self.slider_font.value())
        
        # 更新顶部栏的透明度
        alpha_hex = f"{self.current_alpha:02x}"
        common_bg = f"#161b22{alpha_hex}"
        self.edit_search.setStyleSheet(f"background-color: {common_bg}; border: 1px solid #30363d; color: #c9d1d9; padding: 6px;")
        self.edit_inner_search.setStyleSheet(f"background-color: {common_bg}; border: 1px solid #30363d; color: #ffd700; font-weight: bold;")
        
        # 头部区域
        self.txt_header.setStyleSheet(f"""
            QTextEdit {{
                border: none; 
                background-color: #0d1117{alpha_hex}; 
                font-family: 'Consolas', 'Microsoft YaHei';
                font-size: {self.current_font_size}px;
            }}
        """)

    def change_font_size(self, size):
        """动态调整字体大小和透明度"""
        self.current_font_size = size
        self.lbl_font_val.setText(f"{size}px")
        
        alpha_hex = f"{self.current_alpha:02x}"
        
        # 列表样式
        self.list_results.setStyleSheet(f"""
            QListWidget {{
                background-color: #0d1117{alpha_hex}; 
                border: 1px solid #30363d; 
                border-radius: 6px;
                color: #c9d1d9; 
                padding: 5px;
                font-size: {size}px;
            }}
            QListWidget::item {{ padding: 5px; }}
            QListWidget::item:selected {{ background-color: #1f6feb; border-radius: 4px; color: white; }}
        """)
        
        # 正文内容样式
        self.txt_content.setStyleSheet(f"""
            QTextEdit {{
                background-color: #0d1117{alpha_hex};
                border: 1px solid #30363d;
                color: #c9d1d9; 
                font-family: 'Consolas', 'Microsoft YaHei', monospace;
                font-size: {size}px;
                line-height: 1.5;
            }}
        """)

        # 头部元数据样式 (复用)
        self.txt_header.setStyleSheet(f"""
            QTextEdit {{
                border: none; 
                background-color: #0d1117{alpha_hex}; 
                font-family: 'Consolas', 'Microsoft YaHei';
                font-size: {size}px;
            }}
        """)

    def setup_shortcuts(self):
        self.shortcut_find = QShortcut(QKeySequence("Ctrl+F"), self)
        self.shortcut_find.activated.connect(self.focus_inner_search)

    def focus_inner_search(self):
        if self.isVisible():
            self.edit_inner_search.setFocus()
            self.edit_inner_search.selectAll()

    def open_file_dialog(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "选择索引文件", "", "JSON Files (*.json);;All Files (*)")
        if file_path:
            self.load_file_content(file_path)

    def refresh_data(self):
        if self.last_loaded_path and os.path.exists(self.last_loaded_path):
            self.edit_search.clear()
            self.load_file_content(self.last_loaded_path)
            self.txt_content.append(f"\n🔄 已刷新数据，显示全部内容。")
        else:
            QMessageBox.warning(self, "无法刷新", "尚未加载文件或文件路径已失效。")

    def load_file_content(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8-sig') as f:
                self.data = json.load(f)
            
            self.all_nodes = []
            if isinstance(self.data, dict):
                structure = self.data.get('structure', [self.data])
            elif isinstance(self.data, list):
                structure = self.data
            else:
                raise ValueError("JSON格式错误")

            self._flatten_structure(structure)
            self.last_loaded_path = file_path
            
            self.txt_content.setText(f"✅ 已加载: {os.path.basename(file_path)}\n📊 节点数: {len(self.all_nodes)}\n")
            self.txt_header.clear()
            self.edit_inner_search.clear()

            self.list_results.clear()
            for node in self.all_nodes:
                self._add_item_to_list(node)
                
        except Exception as e:
            import traceback
            self.txt_content.setText(f"❌ 加载失败: {str(e)}\n\n{traceback.format_exc()}")

    def _flatten_structure(self, structure):
        if not structure: return
        for item in structure:
            self.all_nodes.append(item)
            if 'nodes' in item and isinstance(item['nodes'], list):
                self._flatten_structure(item['nodes'])

    def search_content(self):
        query = self.edit_search.text().strip().lower()
        self.list_results.clear()
        
        if not query:
            for node in self.all_nodes: self._add_item_to_list(node)
            return
            
        count = 0
        for node in self.all_nodes:
            if query in node.get('title', '').lower() or query in node.get('text', '').lower():
                self._add_item_to_list(node)
                count += 1
        
        self.txt_content.setText(f"🔍 搜索: '{query}'\n✅ 找到 {count} 个结果。")

    def _add_item_to_list(self, node):
        title = node.get('title', '无标题')
        display = (title[:40] + '...') if len(title) > 40 else title
        item = QListWidgetItem(display)
        item.setToolTip(title)
        item.setData(Qt.UserRole, node)
        self.list_results.addItem(item)

    def display_node_detail(self, item):
        node = item.data(Qt.UserRole)
        if node:
            start, end = node.get('start_index', '-'), node.get('end_index', '-')
            header = f"""
            <h2 style='color: #58a6ff;'>{node.get('title', '未命名')}</h2>
            <div style='background-color: transparent; padding: 5px; color: #c9d1d9;'>
                <b>📄 页码:</b> {start}-{end} &nbsp;|&nbsp; <b>ID:</b> {node.get('node_id', 'N/A')}
            </div>
            """
            self.txt_header.setHtml(header)
            self.txt_content.setPlainText(node.get('text', '无正文'))
            if self.edit_inner_search.text(): self.highlight_text_in_detail()

    def highlight_text_in_detail(self):
        s = self.edit_inner_search.text()
        cursor = self.txt_content.textCursor()
        cursor.select(QTextCursor.Document)
        fmt = QTextCharFormat()
        fmt.setBackground(Qt.transparent)
        cursor.setCharFormat(fmt)
        
        if not s: return
        
        highlight = QTextCharFormat()
        highlight.setBackground(QColor("#d29922"))
        highlight.setForeground(QColor("black"))
        
        cursor.setPosition(0)
        while True:
            cursor = self.txt_content.document().find(s, cursor)
            if cursor.isNull(): break
            cursor.mergeCharFormat(highlight)

    def export_data(self):
        if not self.all_nodes: return
        fmt = self.combo_export.currentText()
        ext = ".docx" if "DOCX" in fmt else ".txt" if "TXT" in fmt else ".csv" if "CSV" in fmt else ".xlsx"
        path, _ = QFileDialog.getSaveFileName(self, "导出", f"export{ext}", f"Files (*{ext})")
        if not path: return

        try:
            if "DOCX" in fmt: self._export_docx(path)
            elif "TXT" in fmt: self._export_txt(path)
            else: self._export_tabular(path, "CSV" in fmt)
            QMessageBox.information(self, "成功", "导出完成")
        except Exception as e:
            QMessageBox.critical(self, "错误", str(e))

    def _export_docx(self, path):
        if not HAS_DOCX: raise ImportError("No python-docx")
        doc = Document()
        for n in self.all_nodes:
            doc.add_heading(n.get('title',''), 1)
            doc.add_paragraph(n.get('text',''))
        doc.save(path)

    def _export_txt(self, path):
        with open(path, 'w', encoding='utf-8') as f:
            for n in self.all_nodes:
                f.write(f"Title: {n.get('title','')}\n{n.get('text','')}\n\n")

    def _export_tabular(self, path, is_csv):
        if not HAS_PANDAS: raise ImportError("No pandas")
        data = [{"Title": n.get('title'), "Text": n.get('text')} for n in self.all_nodes]
        df = pd.DataFrame(data)
        if is_csv: df.to_csv(path, index=False, encoding='utf-8-sig')
        else: df.to_excel(path, index=False)

# =================================================================================
# 核心主窗口: UnifiedMainWindow
# =================================================================================

class UnifiedMainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("PageIndex Pro - Integrated Suite (DeepSeek适配版)")
        self.resize(1200, 900)
        
        self.setStyleSheet(GLOBAL_STYLESHEET)
        
        # === 粒子特效覆盖层 (关键点) ===
        # 将其作为父窗口的子项，并放置在最顶层，但它是“鼠标穿透”的。
        # 当Tab页面的背景透明度降低时，它们会“浮”在半透明背景之上/之中。
        self.overlay = NeuralParticleOverlay(self)
        self.overlay.raise_() 

        self.tabs = QTabWidget()
        self.setCentralWidget(self.tabs)
        
        self.tab_indexer = IndexerTab()
        self.tab_recall = RecallTab()
        
        self.tabs.addTab(self.tab_indexer, "🔧 索引构建 (Indexer)")
        self.tabs.addTab(self.tab_recall, "🔎 知识召回 (Recall)")
        
        tab_bar = self.tabs.tabBar()
        font = tab_bar.font()
        font.setPointSize(11)
        font.setBold(True)
        tab_bar.setFont(font)

    def resizeEvent(self, event):
        self.overlay.resize(self.size())
        super().resizeEvent(event)

if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    window = UnifiedMainWindow()
    window.show()
    sys.exit(app.exec_())
